"""
Generates Medify B.Tech project report as .docx
Run: python scripts/generate_medify_report_docx.py
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill_hex: str):
    """Light gray header shading."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_title_block(doc, lines, bold_center=True):
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.bold = bold_center
        if i == 0:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)


def main():
    out_dir = Path(__file__).resolve().parent.parent
    out_path = out_dir / "Medify_Project_Report.docx"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # --- Cover ---
    doc.add_paragraph()
    doc.add_paragraph()
    add_title_block(
        doc,
        [
            "A Project Report Submitted in Partial Fulfilment of the Requirements",
            "for the Degree of",
            "Bachelor of Technology in",
            "Computer Science and Engineering",
        ],
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Submitted by:")
    r.bold = True
    doc.add_paragraph("[Student Name 1]  ([Roll No.])", style="Normal").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    doc.add_paragraph("[Student Name 2]  ([Roll No.])", style="Normal").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    doc.add_paragraph("(Add all group members)", style="Normal").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Under the Guidance of")
    r.bold = True
    doc.add_paragraph("[Prof. Name]", style="Normal").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("(Designation)", style="Normal").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Department of Computer Science and Engineering", style="Normal"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Sanjivani University, Kopargaon", style="Normal").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    doc.add_paragraph("A. Nagar, Maharashtra – 423601, India", style="Normal").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    doc.add_paragraph()
    doc.add_paragraph("Session [2025–26]", style="Normal").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph("Department of Computer Science and Engineering", style="Normal").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    doc.add_paragraph("Sanjivani University", style="Normal").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # --- Declaration ---
    doc.add_heading("DECLARATION", level=1)
    doc.add_paragraph(
        'We hereby declare that this project entitled "Medify – Healthcare Web Platform" '
        "submitted to the Department of Computer Science and Engineering, Sanjivani University, "
        "Kopargaon, Maharashtra, India for partial fulfilment of the requirement for the degree of "
        "Bachelor of Technology in Computer Science and Engineering is prepared by us and the same "
        "has not been submitted to any other institution."
    )
    doc.add_paragraph()
    doc.add_paragraph("Signatures & names: _________________  _________________")
    doc.add_paragraph("Department of CSE, Sanjivani University")

    doc.add_page_break()

    # --- Certificate Guide ---
    doc.add_heading("CERTIFICATE (Guide)", level=1)
    doc.add_paragraph(
        'This is to certify that the project entitled "Medify – Healthcare Web Platform" submitted '
        "to the Department of Computer Science and Engineering, Sanjivani University, for the major "
        "project of Bachelor of Technology in Computer Science and Engineering is a record of work "
        "carried out by [names] under my supervision and guidance. All help received from various "
        "sources has been duly acknowledged."
    )
    doc.add_paragraph()
    doc.add_paragraph("Prof. [Guide Name]")
    doc.add_paragraph("Professor, Department of CSE")
    doc.add_paragraph("Sanjivani University")
    doc.add_paragraph("Date: __________    Place: Kopargaon")

    doc.add_page_break()

    # --- Certificate HOD ---
    doc.add_heading("CERTIFICATE (HOD)", level=1)
    doc.add_paragraph(
        'This is to certify that the major project entitled "Medify – Healthcare Web Platform" '
        "submitted by [names] is found worthy for the major project of Bachelor of Technology in "
        "Computer Science and Engineering. They have worked under the supervision of Prof. [Guide Name]."
    )
    doc.add_paragraph()
    doc.add_paragraph("Prof. [HOD Name]")
    doc.add_paragraph("(HOD), Department of CSE")
    doc.add_paragraph("Sanjivani University")

    doc.add_page_break()

    # --- Acknowledgement ---
    doc.add_heading("ACKNOWLEDGEMENT", level=1)
    doc.add_paragraph(
        "We express sincere gratitude to our Head of Department, [HOD Name], Sanjivani University, "
        "Kopargaon, and the Department of Computer Science and Engineering for the opportunity and "
        'resources to undertake this project titled "Medify – Healthcare Web Platform". We thank our '
        "project guide, [Guide Name], [Designation], Department of Computer Science and Engineering, "
        "for constant guidance, suggestions, and encouragement. We are grateful to faculty members "
        "and peers for support and feedback during development and testing. This project reflects "
        "collective effort and teamwork."
    )

    doc.add_page_break()

    # --- Abstract ---
    doc.add_heading("ABSTRACT", level=1)
    doc.add_paragraph(
        "Medify is a web-based healthcare platform built with HTML, CSS, and JavaScript, integrated "
        "with Google Firebase for authentication, database, and file storage. It brings together online "
        "medicine ordering, doctor consultation discovery and booking flows, lab test information, "
        "health insurance plan exploration, and secure video consultation in one responsive interface. "
        "The system supports multiple user roles through dedicated flows and dashboards (for example "
        "customers, doctors, pharmacists, delivery, and clinic staff), with Firebase Authentication "
        "and Firestore (and Storage where applicable) for secure, scalable data handling."
    )
    doc.add_paragraph(
        "The platform is designed for accessibility on desktop and mobile browsers, with a clean UI, "
        "shopping cart behaviour, prescription upload for medicine orders, and role-based account "
        "creation and login. Medify demonstrates how modern web technologies plus a Backend-as-a-Service "
        "(BaaS) can deliver a modular, extensible healthcare experience without a traditional custom "
        "server for core features, while remaining suitable for future integration with payment gateways, "
        "appointment systems, and regulatory compliance enhancements."
    )
    p = doc.add_paragraph()
    p.add_run("Keywords: ").bold = True
    p.add_run(
        "Healthcare web application, e-pharmacy, teleconsultation, Firebase, Firestore, responsive web design."
    )

    doc.add_page_break()

    # --- Contents table ---
    doc.add_heading("CONTENTS", level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Sr. No."
    hdr[1].text = "Content"
    hdr[2].text = "Page No."
    for c in hdr:
        set_cell_shading(c, "D9D9D9")
        for para in c.paragraphs:
            for run in para.runs:
                run.bold = True
    rows = [
        ("1", "Introduction", "1"),
        ("2", "Literature Survey", "2"),
        ("3", "Scope of Project", "3"),
        ("4", "Methodology", "4"),
        ("5", "System Design", "5"),
        ("6", "Data Flow Diagram", "6"),
        ("7", "Software Description", "7–9"),
        ("8", "Results and Application", "10–13"),
        ("9", "Conclusion", "14"),
        ("10", "Future Work", "15"),
        ("11", "References", "16"),
    ]
    for sn, content, page in rows:
        row = table.add_row().cells
        row[0].text = sn
        row[1].text = content
        row[2].text = page
    doc.add_paragraph("(Update page numbers after final pagination in Word.)")

    doc.add_page_break()

    # --- I Introduction ---
    doc.add_heading("I. INTRODUCTION", level=1)
    doc.add_paragraph(
        "Healthcare delivery is increasingly digital: patients expect to order medicines, find doctors, "
        "book diagnostics, and understand insurance options from a single digital experience. Medify "
        "addresses this by providing a browser-based platform that unifies these services behind one brand "
        "and one navigation model."
    )
    doc.add_paragraph(
        "Traditionally, users switch between pharmacy websites, hospital portals, lab chains, and insurance "
        "aggregators, which is fragmented and time-consuming. Medify reduces this fragmentation by offering "
        "integrated modules—medicines, consultations, lab tests, insurance, and video consultation—with "
        "consistent design and shared authentication and data layers (Firebase)."
    )
    doc.add_paragraph(
        "The system is implemented as a static front-end (HTML/CSS/JavaScript) with Firebase for user "
        "sign-in, profile and role-related data in Firestore, and Storage for assets such as prescriptions "
        "where implemented. Separate HTML pages and dashboards reflect role-specific workflows (e.g. doctor "
        "vs pharmacist vs delivery), improving clarity and maintainability."
    )
    doc.add_paragraph(
        "The project demonstrates practical software engineering: modular pages, reusable scripts, security "
        "rules on the backend (Firestore/Storage), and a responsive layout for varied devices—aligned with "
        "digital inclusion and usability goals for a university major project."
    )

    doc.add_heading("II. LITERATURE SURVEY", level=1)
    doc.add_paragraph(
        "A literature survey for a digital healthcare platform typically covers: (1) E-health and telemedicine—"
        "research shows that web and mobile health apps improve access when they combine simple navigation, "
        "trust signals, and clear data privacy messaging. (2) E-pharmacy and prescription handling—studies "
        "highlight cart workflows, upload of prescriptions, and verification as critical. (3) User interface "
        "for health apps—literature recommends large touch targets, readable typography, and consistent "
        "terminology. (4) Cloud backends—Firebase is widely used for authentication, real-time sync, and "
        "storage. (5) Security—least-privilege access and database rules; never trusting the client alone."
    )
    doc.add_paragraph(
        "This survey supports Medify’s choices: web-first delivery, Firebase for backend services, and "
        "role-based separation of features."
    )

    doc.add_heading("III. SCOPE OF PROJECT", level=1)
    doc.add_paragraph("End-user (patient/customer) scope:")
    for t in [
        "Browse and order medicines with categories, search, and cart.",
        "Upload prescription for validated ordering workflows.",
        "Consult doctors: specialties, filters (mode, experience, fees, language), sorting.",
        "Lab tests: packages, categories, booking-oriented information.",
        "Health insurance: compare/explore plans.",
        "Video consultation page for remote visits.",
        "Authentication: login/signup, address/delivery context, notifications (as implemented).",
    ]:
        doc.add_paragraph(t, style="List Bullet")
    doc.add_paragraph("Operational / role-based scope:")
    for t in [
        "Doctor dashboard and account creation flows.",
        "Pharmacist dashboard for pharmacy-side operations.",
        "Delivery dashboard for delivery workflows.",
        "Clinic dashboard and clinic account creation where applicable.",
    ]:
        doc.add_paragraph(t, style="List Bullet")
    doc.add_paragraph(
        "Out of scope (typical for academic v1): certified clinical advice beyond demo UI; full production "
        "payment gateway; formal regulatory compliance certification; native mobile apps unless added later."
    )

    doc.add_heading("IV. METHODOLOGY", level=1)
    for t in [
        "Requirements analysis: identify actors (customer, doctor, pharmacist, delivery, clinic); map pages.",
        "Technology selection: HTML5, CSS3, JavaScript; Firebase (Auth, Firestore, Storage); static hosting.",
        "Design: wireframe journeys—browse → cart → checkout; search doctor → profile → consult; login → dashboard.",
        "Implementation: modular HTML, shared styles.css and script.js, feature-specific JS and data files.",
        "Integration: firebase-config.js; Firestore collections for users, orders, roles (document actual schema).",
        "Testing: manual tests on Chrome/Edge, responsive view, auth flows, cart, dashboard access.",
        "Security: document Firestore and Storage rules as implemented.",
    ]:
        doc.add_paragraph(t, style="List Number")

    doc.add_heading("V. SYSTEM DESIGN", level=1)
    doc.add_paragraph(
        "Target audience: patients and families; healthcare partners using role dashboards. "
        "High-level architecture: presentation (HTML/CSS/JS), application logic (client-side JS), "
        "data layer (Firestore, Storage, Auth). Major modules: Pharmacy, Consultation, Lab tests, Insurance, "
        "Video consultation, Auth & profiles. Design principles: responsive layout, consistent header/footer, "
        "clear calls-to-action."
    )

    doc.add_heading("VI. DATA FLOW DIAGRAM", level=1)
    doc.add_paragraph(
        "Level 0 (Context): External entities include User (Customer), Doctor, Pharmacist, Delivery staff, "
        "Clinic staff. The system is the Medify Web Application. Data flows include login credentials, profile "
        "data, cart and order data, prescription files, consultation requests, and dashboard updates. "
        "Logical data stores: Firestore (users, orders, listings), Storage (prescriptions, images)."
    )
    doc.add_paragraph(
        "Level 1 (example processes): Authenticate user ↔ Firebase Auth; Browse catalogue; Manage cart; "
        "Place order → Firestore + Storage; Role dashboard → read/write per security rules. "
        "Include a diagram in the report: User → Medify → Firebase with arrows labeled Auth, Read/Write, File upload."
    )

    doc.add_heading("VII. SOFTWARE DESCRIPTION", level=1)
    doc.add_paragraph(
        "Frontend: HTML5, CSS3, JavaScript for structure, styling, and behaviour (cart, modals, filters, Firebase SDK)."
    )
    doc.add_paragraph(
        "Firebase: Authentication for identity; Cloud Firestore for structured documents; Cloud Storage for "
        "uploads; security via Firestore and Storage rules."
    )
    doc.add_paragraph(
        "Tools: VS Code or Cursor, Git, browser DevTools; optional Firebase CLI for deployment. "
        "Unlike an Android-only stack (e.g. Java + XML), Medify uses web standards and Firebase for cross-platform browser access."
    )

    doc.add_heading("VIII. RESULTS AND APPLICATION", level=1)
    doc.add_paragraph(
        "Results: Home page presents navigation to all modules; Medicines demonstrates search, categories, cart, "
        "prescription upload; Doctors shows filtering and profiles; Lab tests and insurance show structured content; "
        "Login/signup and role dashboards show multi-role design; Video consultation supports remote care presentation."
    )
    doc.add_paragraph(
        "Screenshots to attach: landing page, medicines with cart, doctors with filters, lab/insurance, login/signup, "
        "one dashboard, video consultation page."
    )
    doc.add_paragraph(
        "Application value: Medify is a prototype for a unified healthcare marketplace, extendable with payments, "
        "appointments, and telemedicine integrations."
    )

    doc.add_heading("IX. CONCLUSION", level=1)
    doc.add_paragraph(
        "Medify successfully demonstrates a unified healthcare web platform combining pharmacy, consultations, "
        "diagnostics information, insurance, and video consultation, backed by Firebase for authentication and data. "
        "The responsive interface and role-based dashboards show how a modular system can be built with standard web "
        "technologies. The project highlights digital access, user-centric navigation, and scalable backend integration."
    )

    doc.add_heading("X. FUTURE WORK", level=1)
    for t in [
        "Payment gateway (UPI, cards) and order tracking.",
        "Appointment scheduling with calendar and reminders.",
        "EHR integration at API/concept level.",
        "Push notifications (FCM).",
        "Multilingual UI.",
        "Progressive Web App (PWA) for installability.",
        "Stronger compliance: encryption, audit logs, consent management.",
        "Native Android/iOS apps sharing the same Firebase backend.",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    doc.add_heading("XI. REFERENCES", level=1)
    refs = [
        "Firebase Documentation — https://firebase.google.com/docs",
        "MDN Web Docs (HTML, CSS, JavaScript) — https://developer.mozilla.org/",
        "Google Firestore Security Rules — https://firebase.google.com/docs/firestore/security/get-started",
        "WHO — Digital health — https://www.who.int/health-topics/digital-health",
        "WCAG 2.1 Quick Reference — https://www.w3.org/WAI/WCAG21/quickref/",
        "Android Studio documentation (for general mobile context) — https://developer.android.com/studio/",
        "Add research papers you have actually cited.",
    ]
    for r in refs:
        doc.add_paragraph(r, style="List Bullet")

    doc.save(out_path)
    print(f"Written: {out_path}")


if __name__ == "__main__":
    main()
